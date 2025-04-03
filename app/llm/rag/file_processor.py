import os
from typing import Dict, List, Optional, Tuple
from uuid import UUID

import docx
import pypdf
import pytesseract
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_openai import OpenAIEmbeddings
from PIL import Image

from app.core.config import settings
from app.core.exceptions import (FileProcessingException,
                                 InvalidFileTypeException)
from app.db.models.user_file import UserFile
from app.db.repositories.user_file_repository import UserFileRepository
from app.schemas.file import FileStatus


class FileProcessor:
    """文件处理器，用于提取文件内容并建立向量索引"""

    def __init__(self, file_repo: UserFileRepository):
        self.file_repo = file_repo

        # 设置嵌入模型
        self.embeddings = OpenAIEmbeddings(
            model=settings.EMBEDDING_MODEL, openai_api_key=settings.OPENAI_API_KEY
        )

        # 设置向量存储
        self.vector_store = None
        if settings.VECTOR_DB_TYPE == "chroma":
            self.vector_store = Chroma(
                persist_directory=settings.CHROMA_DB_DIR,
                embedding_function=self.embeddings,
            )

        # 设置文本分割器
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
        )

    async def process_file(self, file_id: UUID) -> Dict:
        """
        处理文件并建立索引

        参数:
            file_id: 文件ID

        返回:
            处理结果信息
        """
        # 获取文件信息
        file_record = await self.file_repo.get_by_id(file_id)
        if not file_record:
            raise FileProcessingException(detail="文件不存在")

        # 更新文件状态为处理中
        await self.file_repo.update_status(
            file_id=file_id, status=FileStatus.PROCESSING
        )

        try:
            # 提取文件内容
            file_path = file_record.storage_path
            content, metadata = self._extract_file_content(
                file_path, file_record.file_type
            )

            # 分割文本
            documents = self._split_text(content, metadata, str(file_id))

            # 创建向量索引
            if self.vector_store:
                self.vector_store.add_documents(documents)
                self.vector_store.persist()

            # 更新文件状态为已索引
            processing_metadata = {
                "chunk_count": len(documents),
                "character_count": len(content),
                **metadata,
            }

            await self.file_repo.update(
                db_obj=file_record,
                obj_in={"status": FileStatus.INDEXED, "metadata": processing_metadata},
            )

            return {
                "status": "success",
                "file_id": str(file_id),
                "metadata": processing_metadata,
            }

        except Exception as e:
            # 更新文件状态为错误
            await self.file_repo.update_status(
                file_id=file_id, status=FileStatus.ERROR, error_message=str(e)
            )

            raise FileProcessingException(detail=f"文件处理失败: {str(e)}")

    def _extract_file_content(self, file_path: str, file_type: str) -> Tuple[str, Dict]:
        """
        提取文件内容

        参数:
            file_path: 文件路径
            file_type: 文件类型

        返回:
            (文件内容, 元数据)
        """
        if not os.path.exists(file_path):
            raise FileProcessingException(detail=f"文件不存在: {file_path}")

        content = ""
        metadata = {}

        try:
            if file_type == "pdf":
                content, metadata = self._extract_pdf(file_path)
            elif file_type == "docx":
                content, metadata = self._extract_docx(file_path)
            elif file_type == "txt":
                content, metadata = self._extract_txt(file_path)
            elif file_type == "image":
                content, metadata = self._extract_image(file_path)
            else:
                raise InvalidFileTypeException(detail=f"不支持的文件类型: {file_type}")

            return content, metadata

        except Exception as e:
            raise FileProcessingException(detail=f"文件内容提取失败: {str(e)}")

    def _extract_pdf(self, file_path: str) -> Tuple[str, Dict]:
        """提取PDF文件内容"""
        with open(file_path, "rb") as file:
            reader = pypdf.PdfReader(file)
            num_pages = len(reader.pages)

            text = ""
            for page_num in range(num_pages):
                page = reader.pages[page_num]
                text += page.extract_text() + "\n\n"

            metadata = {"page_count": num_pages, "source_type": "pdf"}

            return text, metadata

    def _extract_docx(self, file_path: str) -> Tuple[str, Dict]:
        """提取DOCX文件内容"""
        doc = docx.Document(file_path)

        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"

        metadata = {"paragraph_count": len(doc.paragraphs), "source_type": "docx"}

        return text, metadata

    def _extract_txt(self, file_path: str) -> Tuple[str, Dict]:
        """提取TXT文件内容"""
        with open(file_path, "r", encoding="utf-8") as file:
            text = file.read()

        metadata = {"source_type": "txt"}

        return text, metadata

    def _extract_image(self, file_path: str) -> Tuple[str, Dict]:
        """使用OCR提取图片文件内容"""
        if pytesseract.which("tesseract") is None:
            raise FileProcessingException(detail="未安装Tesseract OCR引擎")

        try:
            # 打开图片
            image = Image.open(file_path)

            # 使用Tesseract进行OCR
            text = pytesseract.image_to_string(image, lang="chi_sim+eng")

            metadata = {
                "source_type": "image",
                "image_size": f"{image.size[0]}x{image.size[1]}",
                "ocr_engine": "tesseract",
            }

            return text, metadata

        except Exception as e:
            raise FileProcessingException(detail=f"OCR处理失败: {str(e)}")

    def _split_text(self, text: str, metadata: Dict, file_id: str) -> List[Document]:
        """
        分割文本为块，用于向量存储

        参数:
            text: 文本内容
            metadata: 文件元数据
            file_id: 文件ID

        返回:
            分割后的文档列表
        """
        # 分割文本
        texts = self.text_splitter.split_text(text)

        # 创建Document对象
        documents = []
        for i, chunk in enumerate(texts):
            doc_metadata = {**metadata, "file_id": file_id, "chunk_id": i}
            doc = Document(page_content=chunk, metadata=doc_metadata)
            documents.append(doc)

        return documents
